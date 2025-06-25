from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .bse6 import metric_weights
from financials.models import (
    Company,
    ExtractedBalanceSheet,
    ExtractedIncomeStatement,
    ExtractedCashFlow,
    DerivedFinancialMetric,
    CreditScore
)


def generate_company_report(company_id):
    try:
        company = Company.objects.get(id=company_id)
    except Company.DoesNotExist:
        return {"error": "Company not found"}

    report = {
        "company_name": company.company_name,
        "company_info": {
            "Company Number": company.company_number,
            "PAN": company.company_pan,
            "Business": company.line_of_business,
            "Website": company.website,
            "Incorporated On": str(company.incorporated_on),
            "Location": company.location,
        },
        "financials": {
            "balance_sheet": {},
            "income_statement": {},
            "cash_flow": {},
            "derived_metrics": {},
            "credit_scores": {}
        }
    }

    def load_metrics(queryset):
        data = {}
        for item in queryset:
            data.setdefault(item.metric, {})[item.year] = item.value
        return data

    report["financials"]["balance_sheet"] = load_metrics(
        ExtractedBalanceSheet.objects.filter(company=company))
    report["financials"]["income_statement"] = load_metrics(
        ExtractedIncomeStatement.objects.filter(company=company))
    report["financials"]["cash_flow"] = load_metrics(
        ExtractedCashFlow.objects.filter(company=company))
    report["financials"]["derived_metrics"] = load_metrics(
        DerivedFinancialMetric.objects.filter(company=company))

    credit_entries = CreditScore.objects.filter(company=company)
    for entry in credit_entries:
        report["financials"]["credit_scores"][entry.year] = {
            "score": entry.score,
            "category": entry.category,
            "risk_bucket": entry.risk_bucket
        }

    return report


def style_header_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def style_data_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            para.space_before = Pt(4)
            para.space_after = Pt(4)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def generate_word_report(company_id):
    data = generate_company_report(company_id)
    if "error" in data:
        return None

    doc = Document()
    doc.add_heading(f"📄 Credit Report: {data['company_name']}", 0)

    # 🏢 Company Info Table
    doc.add_heading("🏢 Company Info", level=1)
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Light Grid'
    for label, value in data['company_info'].items():
        row = info_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        style_data_row(info_table.rows[-1])

    # 📊 Credit Scores Table
    doc.add_heading("📊 Credit Score Summary", level=1)
    cs_table = doc.add_table(rows=1, cols=4)
    cs_table.style = 'Colorful List Accent 1'
    hdr = cs_table.rows[0].cells
    hdr[0].text = "Year"
    hdr[1].text = "Score"
    hdr[2].text = "Category"
    hdr[3].text = "Risk Bucket"
    style_header_row(cs_table.rows[0])

    for year, cs in sorted(data["financials"]["credit_scores"].items()):
        row = cs_table.add_row().cells
        row[0].text = year
        row[1].text = f"{cs['score']:.2f}"
        row[2].text = cs['category']
        row[3].text = cs['risk_bucket']
        style_data_row(cs_table.rows[-1])

    # 📈 Derived + Extracted Metrics
    add_metric_section(doc, "📈 Derived Financial Metrics", data["financials"]["derived_metrics"])
    add_metric_section(doc, "📂 Balance Sheet", data["financials"]["balance_sheet"])
    add_metric_section(doc, "📂 Income Statement", data["financials"]["income_statement"])
    add_metric_section(doc, "📂 Cash Flow Statement", data["financials"]["cash_flow"])

    # 📌 Metric Weights
    doc.add_heading("📌 Scoring Metric Weights", level=1)
    wt_table = doc.add_table(rows=1, cols=2)
    wt_table.style = 'Light List'
    wt_header = wt_table.rows[0].cells
    wt_header[0].text = "Metric"
    wt_header[1].text = "Weight (%)"
    style_header_row(wt_table.rows[0])

    for metric, wt in metric_weights.items():
        row = wt_table.add_row().cells
        row[0].text = metric
        row[1].text = f"{wt * 100:.1f}"
        style_data_row(wt_table.rows[-1])

    # Final Output
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_metric_section(doc, title, metric_data):
    doc.add_heading(title, level=1)
    if not metric_data:
        doc.add_paragraph("No data available.")
        return

    metrics = sorted(metric_data.keys())
    years = sorted({year for values in metric_data.values() for year in values})

    table = doc.add_table(rows=1, cols=len(years) + 1)
    table.style = 'Colorful List Accent 1'
    header = table.rows[0].cells
    header[0].text = "Metric"
    for i, y in enumerate(years):
        header[i + 1].text = y
    style_header_row(table.rows[0])

    for metric in metrics:
        row = table.add_row().cells
        row[0].text = metric
        for i, y in enumerate(years):
            val = metric_data[metric].get(y, "-")
            row[i + 1].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)
        style_data_row(table.rows[-1])
